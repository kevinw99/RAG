"""Document processing pipeline with multi-format support.

Optimized for 612MB document library (1,566 documents) with efficient chunking
and memory management.
"""

import hashlib
import logging
import mimetypes
import time
from pathlib import Path
from typing import List, Dict, Optional, Generator, Tuple, Any
from datetime import datetime
import asyncio
import gc

# Document format processors
import PyPDF2
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import markdown

from ..core.data_models import Document, Chunk, DocumentType, ProcessingStats
from ..utils.text_processing import text_processor
from ..config.settings import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Document processor with multi-format support and semantic chunking.
    
    Optimized for large datasets with memory management and batch processing.
    """
    
    def __init__(self, 
                 chunk_size: int = None,
                 chunk_overlap: int = None,
                 batch_size: int = None):
        """Initialize document processor.
        
        Args:
            chunk_size: Size of text chunks (default from settings)
            chunk_overlap: Overlap between chunks (default from settings)
            batch_size: Batch size for processing (default from settings)
        """
        self.chunk_size = chunk_size or settings.chunk_size
        self.chunk_overlap = chunk_overlap or settings.chunk_overlap
        self.batch_size = batch_size or settings.batch_size
        
        self.stats = ProcessingStats()
        
        # Document processors for different formats
        self.processors = {
            DocumentType.PDF: self._process_pdf,
            DocumentType.TEXT: self._process_text,
            DocumentType.MARKDOWN: self._process_markdown,
            DocumentType.DOCX: self._process_docx,
            DocumentType.DOC: self._process_doc,
            DocumentType.HTML: self._process_html,
        }
        
        logger.info(f"DocumentProcessor initialized: chunk_size={self.chunk_size}, "
                   f"chunk_overlap={self.chunk_overlap}, batch_size={self.batch_size}")
    
    async def process_directory(self, directory: Path, 
                               recursive: bool = True) -> List[Document]:
        """Process all documents in a directory.
        
        Args:
            directory: Directory containing documents
            recursive: Whether to process subdirectories
            
        Returns:
            List of processed documents
        """
        logger.info(f"Processing directory: {directory}")
        
        if not directory.exists():
            raise ValueError(f"Directory does not exist: {directory}")
        
        # Find all supported files
        files = self._find_supported_files(directory, recursive)
        logger.info(f"Found {len(files)} supported files")
        
        self.stats.total_documents = len(files)
        
        # Process files in batches for memory efficiency
        documents = []
        for i in range(0, len(files), self.batch_size):
            batch_files = files[i:i + self.batch_size]
            
            logger.info(f"Processing batch {i//self.batch_size + 1}/{(len(files)-1)//self.batch_size + 1}")
            
            batch_docs = await self._process_file_batch(batch_files)
            documents.extend(batch_docs)
            
            # Force garbage collection for memory management
            if settings.enable_garbage_collection:
                gc.collect()
        
        self.stats.processing_time = time.time() - start_time if 'start_time' in locals() else 0
        
        logger.info(f"Processing complete: {self.stats.processed_documents}/{self.stats.total_documents} "
                   f"documents, {self.stats.total_chunks} chunks")
        
        return documents
    
    def _find_supported_files(self, directory: Path, recursive: bool) -> List[Path]:
        """Find all supported files in directory."""
        files = []
        
        pattern = "**/*" if recursive else "*"
        for file_path in directory.glob(pattern):
            if not file_path.is_file():
                continue
                
            # Check file extension
            if file_path.suffix.lower() in settings.supported_extensions:
                # Check file size
                file_size = file_path.stat().st_size
                if file_size > settings.max_file_size_mb * 1024 * 1024:
                    logger.warning(f"Skipping large file: {file_path} ({file_size/1024/1024:.1f}MB)")
                    continue
                    
                files.append(file_path)
                self.stats.file_sizes[str(file_path)] = file_size
        
        return files
    
    async def _process_file_batch(self, file_paths: List[Path]) -> List[Document]:
        """Process a batch of files concurrently."""
        tasks = [self.process_file(file_path) for file_path in file_paths]
        
        # Process with limited concurrency to manage memory
        semaphore = asyncio.Semaphore(5)  # Max 5 concurrent file processes
        
        async def process_with_semaphore(task):
            async with semaphore:
                return await task
        
        results = await asyncio.gather(
            *[process_with_semaphore(task) for task in tasks],
            return_exceptions=True
        )
        
        documents = []
        for result in results:
            if isinstance(result, Document):
                documents.append(result)
                self.stats.processed_documents += 1
            elif isinstance(result, Exception):
                self.stats.errors.append(str(result))
                logger.error(f"Error processing file: {result}")
        
        return documents
    
    async def process_file(self, file_path: Path) -> Document:
        """Process a single file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Processed document
        """
        logger.debug(f"Processing file: {file_path}")
        
        try:
            # Determine document type
            doc_type = self._get_document_type(file_path)
            
            # Generate document ID
            doc_id = self._generate_doc_id(file_path)
            
            # Process based on type
            content = await self.processors[doc_type](file_path)
            
            # Clean content
            content = text_processor.clean_text(content)
            
            if not content.strip():
                raise ValueError("No content extracted from document")
            
            # Extract metadata
            metadata = self._extract_file_metadata(file_path, content)
            
            document = Document(
                content=content,
                metadata=metadata,
                doc_id=doc_id,
                source=file_path,
                doc_type=doc_type,
                file_size=file_path.stat().st_size
            )
            
            logger.debug(f"Successfully processed: {file_path} ({len(content)} chars)")
            return document
            
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            raise
    
    def _get_document_type(self, file_path: Path) -> DocumentType:
        """Determine document type from file extension."""
        suffix = file_path.suffix.lower()
        
        type_mapping = {
            '.pdf': DocumentType.PDF,
            '.txt': DocumentType.TEXT,
            '.md': DocumentType.MARKDOWN,
            '.docx': DocumentType.DOCX,
            '.doc': DocumentType.DOC,
            '.html': DocumentType.HTML,
            '.htm': DocumentType.HTML,
        }
        
        doc_type = type_mapping.get(suffix)
        if not doc_type:
            raise ValueError(f"Unsupported file type: {suffix}")
        
        return doc_type
    
    def _generate_doc_id(self, file_path: Path) -> str:
        """Generate unique document ID."""
        # Use file path and modification time for uniqueness
        content = f"{file_path}_{file_path.stat().st_mtime}"
        return hashlib.sha256(content.encode()).hexdigest()[:16]
    
    async def _process_pdf(self, file_path: Path) -> str:
        """Process PDF file."""
        try:
            content = []
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            content.append(text)
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num} from {file_path}: {e}")
            
            return '\n\n'.join(content)
            
        except Exception as e:
            raise ValueError(f"Error processing PDF {file_path}: {e}")
    
    async def _process_text(self, file_path: Path) -> str:
        """Process plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        except Exception as e:
            raise ValueError(f"Error processing text file {file_path}: {e}")
    
    async def _process_markdown(self, file_path: Path) -> str:
        """Process Markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                md_content = file.read()
            
            # Convert markdown to HTML, then extract text
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text(separator='\n', strip=True)
            
        except Exception as e:
            raise ValueError(f"Error processing markdown file {file_path}: {e}")
    
    async def _process_docx(self, file_path: Path) -> str:
        """Process DOCX file."""
        try:
            doc = DocxDocument(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            return '\n\n'.join(content)
            
        except Exception as e:
            raise ValueError(f"Error processing DOCX file {file_path}: {e}")
    
    async def _process_doc(self, file_path: Path) -> str:
        """Process DOC file using multiple fallback approaches."""
        try:
            # Method 1: Try to read as DOCX (some .doc files are actually DOCX format)
            try:
                doc = DocxDocument(file_path)
                content = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content.append(paragraph.text)
                
                if content:
                    return '\n\n'.join(content)
                    
            except Exception:
                pass
            
            # Method 2: Try using docx2txt (handles some DOC files)
            try:
                import docx2txt
                content = docx2txt.process(str(file_path))
                if content and content.strip():
                    return content.strip()
            except Exception:
                pass
            
            # Method 3: Try antiword if available (Linux/Mac tool for DOC files)
            try:
                import subprocess
                result = subprocess.run(['antiword', str(file_path)], 
                                      capture_output=True, text=True, timeout=30)
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout.strip()
            except (FileNotFoundError, subprocess.TimeoutExpired, Exception):
                pass
            
            # Method 4: Try catdoc if available (another DOC converter)
            try:
                import subprocess
                result = subprocess.run(['catdoc', str(file_path)], 
                                      capture_output=True, text=True, timeout=30)
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout.strip()
            except (FileNotFoundError, subprocess.TimeoutExpired, Exception):
                pass
            
            # Method 5: Last resort - extract readable text from binary content
            try:
                with open(file_path, 'rb') as file:
                    binary_content = file.read()
                
                # Try to decode as various encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252', 'utf-16']:
                    try:
                        text_content = binary_content.decode(encoding, errors='ignore')
                        # Clean up binary artifacts and extract readable text
                        import re
                        # Remove control characters but keep newlines and tabs
                        clean_text = re.sub(r'[^\x20-\x7E\n\r\t\u00A0-\uFFFF]', ' ', text_content)
                        # Remove excessive whitespace
                        clean_text = re.sub(r'\s+', ' ', clean_text)
                        # Extract words that look like real text (at least 3 chars, mostly letters)
                        words = []
                        for word in clean_text.split():
                            if len(word) >= 3 and sum(c.isalpha() for c in word) >= len(word) * 0.7:
                                words.append(word)
                        
                        if len(words) > 10:  # If we found reasonable amount of text
                            return ' '.join(words)
                            
                    except Exception:
                        continue
            except Exception:
                pass
            
            # If all methods fail, return empty string (will be handled as empty content)
            return ""
            
        except Exception as e:
            logger.warning(f"Error processing DOC file {file_path}: {e}")
            return ""  # Return empty instead of raising error
    
    async def _process_html(self, file_path: Path) -> str:
        """Process HTML file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            return soup.get_text(separator='\n', strip=True)
            
        except Exception as e:
            raise ValueError(f"Error processing HTML file {file_path}: {e}")
    
    def _extract_file_metadata(self, file_path: Path, content: str) -> Dict[str, Any]:
        """Extract metadata from file and content."""
        stat = file_path.stat()
        
        # Basic file metadata
        metadata = {
            'filename': file_path.name,
            'file_path': str(file_path),
            'file_size': stat.st_size,
            'created_at': datetime.fromtimestamp(stat.st_ctime),
            'modified_at': datetime.fromtimestamp(stat.st_mtime),
            'file_extension': file_path.suffix.lower(),
        }
        
        # Content metadata
        content_metadata = text_processor.extract_metadata_from_text(content)
        metadata.update(content_metadata)
        
        # MIME type
        mime_type, _ = mimetypes.guess_type(file_path)
        if mime_type:
            metadata['mime_type'] = mime_type
        
        return metadata
    
    def create_chunks(self, document: Document, 
                     enable_contextual: bool = True) -> List[Chunk]:
        """Create chunks from document with semantic awareness.
        
        Args:
            document: Document to chunk
            enable_contextual: Enable contextual retrieval preprocessing
            
        Returns:
            List of document chunks
        """
        if not document.content:
            return []
        
        chunks = []
        content = document.content
        
        # Create document context for contextual retrieval
        doc_context = f"{document.metadata.get('filename', 'Unknown')} - " \
                     f"{content[:100]}..." if enable_contextual else ""
        
        # Split into chunks with overlap
        start_pos = 0
        chunk_index = 0
        
        while start_pos < len(content):
            # Calculate end position
            end_pos = min(start_pos + self.chunk_size, len(content))
            
            # Try to break at sentence boundary if not at end of document
            if end_pos < len(content):
                # Look for sentence endings in the last 20% of chunk
                search_start = int(end_pos - self.chunk_size * 0.2)
                sentence_end = self._find_sentence_boundary(content, search_start, end_pos)
                if sentence_end > start_pos:
                    end_pos = sentence_end
            
            # Extract chunk content
            chunk_content = content[start_pos:end_pos].strip()
            
            if not chunk_content:
                break
            
            # Apply contextual retrieval if enabled
            if enable_contextual and doc_context:
                chunk_content = text_processor.create_contextual_chunk(
                    chunk_content, doc_context
                )
            
            # Create chunk
            chunk_id = f"{document.doc_id}_{chunk_index:04d}"
            
            chunk = Chunk(
                content=chunk_content,
                metadata={
                    **document.metadata,
                    'chunk_index': chunk_index,
                    'contextual': enable_contextual,
                },
                chunk_id=chunk_id,
                doc_id=document.doc_id,
                start_char=start_pos,
                end_char=end_pos,
                chunk_index=chunk_index
            )
            
            chunks.append(chunk)
            chunk_index += 1
            
            # Move start position with overlap
            if end_pos >= len(content):
                break
                
            start_pos = end_pos - self.chunk_overlap
            
            # Ensure we make progress
            if start_pos <= chunks[-1].start_char:
                start_pos = end_pos
        
        self.stats.total_chunks += len(chunks)
        
        logger.debug(f"Created {len(chunks)} chunks for document {document.doc_id}")
        return chunks
    
    def _find_sentence_boundary(self, text: str, start: int, end: int) -> int:
        """Find sentence boundary for better chunk breaks."""
        sentence_endings = '.!?'
        
        # Look backwards from end position
        for i in range(end - 1, start - 1, -1):
            if text[i] in sentence_endings:
                # Make sure it's not an abbreviation (simple check)
                if i < len(text) - 1 and text[i + 1].isspace():
                    return i + 1
        
        return end
    
    def get_processing_stats(self) -> ProcessingStats:
        """Get processing statistics."""
        return self.stats
    
    def reset_stats(self) -> None:
        """Reset processing statistics."""
        self.stats = ProcessingStats()


# Convenience function for quick processing
async def process_documents(directory: Path, **kwargs) -> Tuple[List[Document], List[Chunk]]:
    """Convenience function to process documents and create chunks.
    
    Returns:
        Tuple of (documents, chunks)
    """
    processor = DocumentProcessor(**kwargs)
    
    # Process documents
    documents = await processor.process_directory(directory)
    
    # Create chunks
    all_chunks = []
    for doc in documents:
        chunks = processor.create_chunks(doc, enable_contextual=True)
        all_chunks.extend(chunks)
    
    return documents, all_chunks